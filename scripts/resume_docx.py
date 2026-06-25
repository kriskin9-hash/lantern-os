#!/usr/bin/env python3
"""
resume_docx.py — docx read/write helper for the Keystone chat document-processing
feature (routes/documents.js). Two modes:

  extract <in.docx>           -> prints the document's plain text to stdout
  build   <out.docx>          -> reads a JSON spec on stdin, writes a styled .docx

The JSON `build` spec is intentionally generic so it works for resumes and other
simple documents:

  {
    "name": "Alexander Place",                 # big title line (optional)
    "contact": "email · phone · City, ST · linkedin.com/in/...",  # subtitle (optional)
    "sections": [
      { "heading": "Professional Summary", "paragraphs": ["..."] },
      { "heading": "Work Experience", "entries": [
          { "title": "Software Engineer", "org": "US Bank — Cincinnati, OH",
            "dates": "Mar 2016 – Apr 2023", "bullets": ["...", "..."] }
      ]},
      { "heading": "Skills", "bullets": ["...", "..."] }
    ]
  }

Each section may carry any of: paragraphs[], bullets[], entries[]. Unknown keys are
ignored, so the LLM can be a little loose without breaking generation.
"""
import json
import sys


def extract(path):
    import docx
    doc = docx.Document(path)
    lines = []
    # Name + contact often live in the page HEADER (not the body) — include them first
    # so the rewrite keeps the person's identity instead of emitting a placeholder.
    seen_hdr = set()
    for section in doc.sections:
        for part in (section.header, section.footer):
            try:
                for p in part.paragraphs:
                    txt = p.text.strip()
                    if txt and txt not in seen_hdr:
                        seen_hdr.add(txt)
                        lines.append(txt)
            except Exception:
                pass
    for p in doc.paragraphs:
        lines.append(p.text)
    # include simple table text too (some resumes use tables for layout)
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append("\t".join(cells))
    sys.stdout.write("\n".join(lines))


def _set_repeat(run):
    run.font.name = "Calibri"


def build(path, spec):
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = docx.Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = docx.shared.Inches(0.6)
        section.left_margin = section.right_margin = docx.shared.Inches(0.7)

    accent = RGBColor(0x1F, 0x4E, 0x79)  # deep blue

    def add_bottom_border(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F4E79")
        pbdr.append(bottom)
        pPr.append(pbdr)

    # Name / title
    name = (spec.get("name") or "").strip()
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name)
        r.bold = True
        r.font.size = Pt(22)
        r.font.color.rgb = accent
        p.space_after = Pt(0)

    contact = (spec.get("contact") or "").strip()
    if contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(contact)
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for sec in spec.get("sections", []):
        heading = (sec.get("heading") or "").strip()
        if heading:
            hp = doc.add_paragraph()
            hp.paragraph_format.space_before = Pt(10)
            hp.paragraph_format.space_after = Pt(2)
            hr = hp.add_run(heading.upper())
            hr.bold = True
            hr.font.size = Pt(11.5)
            hr.font.color.rgb = accent
            add_bottom_border(hp)

        for para in sec.get("paragraphs", []) or []:
            if not str(para).strip():
                continue
            pp = doc.add_paragraph(str(para).strip())
            pp.paragraph_format.space_after = Pt(4)

        for entry in sec.get("entries", []) or []:
            title = (entry.get("title") or "").strip()
            org = (entry.get("org") or "").strip()
            dates = (entry.get("dates") or "").strip()
            line = doc.add_paragraph()
            line.paragraph_format.space_before = Pt(4)
            line.paragraph_format.space_after = Pt(0)
            left = " — ".join([x for x in [title, org] if x])
            lr = line.add_run(left)
            lr.bold = True
            lr.font.size = Pt(10.5)
            if dates:
                # right-aligned dates via tab stop
                from docx.enum.text import WD_TAB_ALIGNMENT
                from docx.shared import Inches as _In
                line.paragraph_format.tab_stops.add_tab_stop(_In(7.1), WD_TAB_ALIGNMENT.RIGHT)
                dr = line.add_run("\t" + dates)
                dr.italic = True
                dr.font.size = Pt(9.5)
                dr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            for b in entry.get("bullets", []) or []:
                if not str(b).strip():
                    continue
                bp = doc.add_paragraph(str(b).strip(), style="List Bullet")
                bp.paragraph_format.space_after = Pt(1)

        for b in sec.get("bullets", []) or []:
            if not str(b).strip():
                continue
            bp = doc.add_paragraph(str(b).strip(), style="List Bullet")
            bp.paragraph_format.space_after = Pt(1)

    doc.save(path)
    sys.stdout.write(json.dumps({"ok": True, "path": path}))


# ── Markdown <-> docx for Document Mode (routes/docmode.js) ──────────────────
def extract_md(path):
    """docx -> markdown: headings (#/##/###), bullets (-), paragraphs."""
    import docx
    doc = docx.Document(path)
    out = []
    seen = set()
    for section in doc.sections:
        for part in (section.header,):
            try:
                for p in part.paragraphs:
                    t = p.text.strip()
                    if t and t not in seen:
                        seen.add(t)
                        out.append(t)
            except Exception:
                pass
    if out:
        out.append("")
    for p in doc.paragraphs:
        t = p.text.rstrip()
        style = (p.style.name or "").lower() if p.style else ""
        if not t.strip():
            out.append("")
            continue
        if "heading 1" in style or style == "title":
            out.append("# " + t.strip())
        elif "heading 2" in style:
            out.append("## " + t.strip())
        elif "heading 3" in style:
            out.append("### " + t.strip())
        elif "list" in style or "bullet" in style:
            out.append("- " + t.strip())
        else:
            out.append(t)
    sys.stdout.write("\n".join(out))


def _add_md_runs(paragraph, text):
    """Render minimal inline markdown (**bold**, *italic*) into runs."""
    import re as _re
    parts = _re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for seg in parts:
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            paragraph.add_run(seg[2:-2]).bold = True
        elif seg.startswith("*") and seg.endswith("*"):
            paragraph.add_run(seg[1:-1]).italic = True
        else:
            paragraph.add_run(seg)


def build_md(path, md):
    """markdown -> styled docx (headings, bullets, paragraphs)."""
    import docx
    from docx.shared import Pt, Inches
    doc = docx.Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.8)
        s.left_margin = s.right_margin = Inches(0.9)

    for raw in (md or "").split("\n"):
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.lstrip().startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_md_runs(p, line.lstrip()[2:].strip())
        else:
            p = doc.add_paragraph()
            _add_md_runs(p, line)
    doc.save(path)
    sys.stdout.write(json.dumps({"ok": True, "path": path}))


def main():
    if len(sys.argv) < 3:
        sys.stderr.write("usage: resume_docx.py <extract|build|extract-md|build-md> <path>\n")
        sys.exit(2)
    mode, path = sys.argv[1], sys.argv[2]
    if mode == "extract":
        extract(path)
    elif mode == "build":
        build(path, json.loads(sys.stdin.read() or "{}"))
    elif mode == "extract-md":
        extract_md(path)
    elif mode == "build-md":
        build_md(path, sys.stdin.read())
    else:
        sys.stderr.write("unknown mode: %s\n" % mode)
        sys.exit(2)


if __name__ == "__main__":
    main()
